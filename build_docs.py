"""Generator dokumentacji technicznej projektu Traffic Light Assistant."""

import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- Paleta kolorow ---
NAVY = RGBColor(0x1F, 0x3A, 0x5F)
BLUE = RGBColor(0x2E, 0x75, 0xB6)
LIGHT_BLUE = RGBColor(0xDE, 0xEB, 0xF7)
MID_BLUE = RGBColor(0x9D, 0xC3, 0xE6)
GREY_TEXT = RGBColor(0x40, 0x40, 0x40)
CODE_BG = "F2F2F2"
HEADER_BG = "1F3A5F"
HEADER_BG_ALT = "2E75B6"

FIG_DIR = r"C:\Users\Patryk\PycharmProjects\LAB04\LAB13\figures"

doc = Document()

# --- Domyslny styl dokumentu ---
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
normal.paragraph_format.line_spacing = 1.15

# --- Marginesy ---
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)


def set_heading_style(name, size, color, bold=True, space_before=18, space_after=8, all_caps=False):
    style = doc.styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(space_before)
    style.paragraph_format.space_after = Pt(space_after)
    style.paragraph_format.keep_with_next = True
    if all_caps:
        rpr = style.element.get_or_add_rPr()
        caps = OxmlElement("w:caps")
        rpr.append(caps)


set_heading_style("Title", 30, NAVY, space_before=0, space_after=6)
set_heading_style("Heading 1", 18, NAVY, space_before=26, space_after=10)
set_heading_style("Heading 2", 14, BLUE, space_before=14, space_after=6)
set_heading_style("Heading 3", 12, NAVY, space_before=10, space_after=4)


def add_heading_border(paragraph, color="2E75B6", sz=8):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def h1(text):
    p = doc.add_heading(text, level=1)
    add_heading_border(p, color="2E75B6", sz=10)
    return p


def h2(text):
    return doc.add_heading(text, level=2)


def h3(text):
    return doc.add_heading(text, level=3)


def para(text, bold=False, italic=False, size=11, align=None, color=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + 0.5 * level)
    return p


def numbered(text):
    return doc.add_paragraph(text, style="List Number")


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=10, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)


def make_table(headers, rows, col_widths=None, header_color=HEADER_BG_ALT):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, htext in enumerate(headers):
        set_cell_text(hdr_cells[i], htext, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(hdr_cells[i], header_color)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val), size=10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def code_block(code_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), CODE_BG)
    pPr.append(shd)
    pbdr = OxmlElement("w:pBdr")
    for side in ("left", "top", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), "9DC3E6")
        pbdr.append(el)
    pPr.append(pbdr)
    lines = code_text.strip("\n").split("\n")
    for i, line in enumerate(lines):
        run = p.add_run(line if line.strip() else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        if i < len(lines) - 1:
            run.add_break()
    return p


def caption(text):
    p = para(text, italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x5A, 0x5A, 0x5A))
    p.paragraph_format.space_before = Pt(2)
    return p


def add_figure(filename, width_cm, cap_text):
    doc.add_picture(f"{FIG_DIR}\\{filename}", width=Cm(width_cm))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption(cap_text)


def page_break():
    doc.add_page_break()


# =====================================================================
# STRONA TYTULOWA
# =====================================================================

for _ in range(4):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("TRAFFIC LIGHT ASSISTANT")
run.font.size = Pt(34)
run.font.bold = True
run.font.color.rgb = NAVY
add_heading_border(title_p, color="2E75B6", sz=18)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run("Dokumentacja techniczna projektu")
run.font.size = Pt(16)
run.font.color.rgb = BLUE
sub_p.paragraph_format.space_before = Pt(10)

sub2_p = doc.add_paragraph()
sub2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2_p.add_run(
    "Wieloagentowy system uczenia ze wzmocnieniem (Multi-Agent Reinforcement Learning)\n"
    "do optymalizacji sygnalizacji swietlnej na skrzyzowaniu dwoch drog"
)
run.font.size = Pt(11)
run.italic = True
run.font.color.rgb = GREY_TEXT
sub2_p.paragraph_format.space_before = Pt(6)

for _ in range(6):
    doc.add_paragraph()

course_p = doc.add_paragraph()
course_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = course_p.add_run("Przedmiot: Analiza regresji i szeregow czasowych")
run.font.size = Pt(13)
run.font.bold = True
run.font.color.rgb = NAVY

for _ in range(4):
    doc.add_paragraph()

authors_p = doc.add_paragraph()
authors_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = authors_p.add_run("Autorzy:")
run.font.size = Pt(12)
run.font.bold = True
run.font.color.rgb = BLUE

for name in ["Patryk Monarcha", "Aleksander Oleszkiewicz", "Julia Labedzka"]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(name)
    r.font.size = Pt(12)

for _ in range(6):
    doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = date_p.add_run(datetime.date.today().strftime("Czerwiec %Y"))
r.font.size = Pt(11)
r.font.color.rgb = GREY_TEXT

page_break()

# =====================================================================
# SPIS TRESCI
# =====================================================================

h1("Spis tresci")

toc_p = doc.add_paragraph()
run = toc_p.add_run()
fldChar = OxmlElement("w:fldChar")
fldChar.set(qn("w:fldCharType"), "begin")
instrText = OxmlElement("w:instrText")
instrText.set(qn("xml:space"), "preserve")
instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
fldChar2 = OxmlElement("w:fldChar")
fldChar2.set(qn("w:fldCharType"), "separate")
fldText = OxmlElement("w:t")
fldText.text = "Aktualizuj pole spisu tresci, aby zobaczyc numery stron (Ctrl+A, F9)."
fldChar3 = OxmlElement("w:t")
fldChar3.text = ""
fldChar4 = OxmlElement("w:fldChar")
fldChar4.set(qn("w:fldCharType"), "end")

r_element = run._r
r_element.append(fldChar)
r_element.append(instrText)
r_element.append(fldChar2)
r_element.append(fldText)
r_element.append(fldChar4)

page_break()

# =====================================================================
# 1. WPROWADZENIE
# =====================================================================

h1("1. Wprowadzenie")

h2("1.1. Cel i kontekst projektu")
para(
    "Projekt „Traffic Light Assistant” jest implementacja systemu wieloagentowego "
    "uczenia ze wzmocnieniem (Multi-Agent Reinforcement Learning, MARL), ktory uczy sie "
    "optymalnego sterowania sygnalizacja swietlna na pojedynczym skrzyzowaniu dwoch drog. "
    "Projekt zostal zrealizowany jako praca semestralna na przedmiot Analiza regresji i "
    "szeregow czasowych i laczy w sobie elementy modelowania procesow stochastycznych "
    "(proces Poissona), uczenia maszynowego (algorytm PPO), inzynierii oprogramowania "
    "(architektura modulowa, REST API, testy automatyczne) oraz wizualizacji danych."
)
para(
    "Glownym celem projektu jest pokazanie, ze agent wytrenowany metoda uczenia ze "
    "wzmocnieniem jest w stanie nauczyc sie adaptacyjnej strategii przelaczania swiatel, "
    "ktora w sposob mierzalny przewyzsza klasyczny, sztywny harmonogram cykliczny "
    "(fixed-cycle baseline). Dodatkowym celem jest zbudowanie kompletnego, dzialajacego "
    "systemu - od symulacji ruchu, przez trening i optymalizacje hiperparametrow, po "
    "API monitorujace i wizualizacje wynikow."
)

h2("1.2. Sformulowanie problemu")
para(
    "Rozwazane jest pojedyncze skrzyzowanie dwoch jednopasmowych drog, oznaczonych jako "
    "droga A i droga B. Pojazdy napotykaja na skrzyzowanie zgodnie z procesem Poissona - "
    "liczba pojazdow przybywajacych w danym kroku czasowym jest zmienna losowa o "
    "rozkladzie Poissona z parametrem λ (lambda). W kazdym kroku symulacji tylko "
    "jedna droga moze posiadac sygnal zielony. Celem ukladu sterujacego jest minimalizacja "
    "calkowitej liczby pojazdow czekajacych na obu drogach, z zachowaniem fizycznych "
    "ograniczen sygnalizacji: minimalnego czasu trwania zielonego swiatla oraz okresu "
    "przejsciowego (zoltego)."
)
para(
    "Problem ten jest klasycznym przykladem decyzyjnym w warunkach niepewnosci - "
    "sterownik nie wie, ile pojazdow przybedzie w nastepnym kroku, a jego decyzje "
    "(przelaczyc / nie przelaczac swiatla) wplywaja na przyszly stan kolejek w sposob "
    "kumulatywny. Jest to zatem naturalne zadanie dla uczenia ze wzmocnieniem, w ktorym "
    "agent maksymalizuje sume nagrod w czasie, a nie nagrode chwilowa."
)

h2("1.3. Zastosowany stos technologiczny")
make_table(
    ["Warstwa / obszar", "Technologia", "Rola w projekcie"],
    [
        ["Srodowisko symulacyjne", "PettingZoo (ParallelEnv)", "Definicja wieloagentowego srodowiska RL zgodnego ze standardem API"],
        ["Adaptacja do SB3", "SuperSuit", "Konwersja srodowiska PettingZoo do wektoryzowanego srodowiska kompatybilnego ze Stable-Baselines3"],
        ["Algorytm uczenia", "Stable-Baselines3 (PPO)", "Trening polityki neuronowej (MlpPolicy) sterujacej obiema sygnalizacjami"],
        ["Obliczenia numeryczne", "NumPy", "Generowanie napływow Poissona, operacje na kolejkach"],
        ["Optymalizacja hiperparametrow", "Optuna", "Przeszukiwanie przestrzeni hiperparametrow PPO (TPE)"],
        ["Stan i metryki w czasie rzeczywistym", "Redis", "Przechowywanie stanu srodowiska, metryk treningu i statusu"],
        ["API", "FastAPI + Uvicorn", "REST API do monitorowania, konfiguracji, treningu i ewaluacji"],
        ["Walidacja danych / konfiguracja", "Pydantic", "Modele konfiguracji (EnvConfig, TrainConfig, RedisConfig)"],
        ["Wizualizacja", "Matplotlib", "Renderowanie skrzyzowania oraz wykresy wynikow"],
        ["Testy", "Pytest, FastAPI TestClient", "Testy jednostkowe i integracyjne"],
        ["Zarzadzanie zaleznosciami", "uv (Python 3.13)", "Instalacja i blokowanie wersji pakietow (uv.lock)"],
        ["Konteneryzacja", "Docker Compose", "Uslugi pomocnicze (Redis)"],
    ],
    col_widths=[4.5, 4.5, 7.5],
)

page_break()

# =====================================================================
# 2. ARCHITEKTURA SYSTEMU
# =====================================================================

h1("2. Architektura systemu")

para(
    "System zostal zaprojektowany w sposob modulowy - kazda warstwa odpowiada za jeden, "
    "dobrze odgraniczony obszar odpowiedzialnosci. Dzieki temu mozliwe jest niezalezne "
    "testowanie logiki symulacji (bez zaleznosci od frameworkow RL), niezalezne "
    "uruchamianie treningu (z lub bez Redis) oraz niezalezne uruchamianie API "
    "monitorujacego."
)

h2("2.1. Diagram warstw systemu")
code_block(
"""
+--------------------------------------------------------------+
|                PettingZoo Environment (2 agentow)            |
|   +----------------+        +----------------+               |
|   |   light_A      |        |   light_B      |               |
|   +-------+--------+        +-------+--------+               |
|           |          SuperSuit       |                        |
|           +-----------+--------------+                        |
|                       v                                        |
|             SB3 VecEnv (wspolna polityka)                      |
+-----------------------+----------------------------------------+
                        v
              Trening PPO (MlpPolicy)
                        |
        +---------------+----------------+
        v                v                v
     Redis           Modele (models/)   FastAPI (api/server.py)
   (metryki)         (checkpointy)        (REST API)
"""
)

h2("2.2. Przeplyw danych podczas treningu")
para("Sekwencja wywolan podczas treningu modelu PPO wyglada nastepujaco:")
numbered("scripts/train.py - wczytanie konfiguracji domyslnej lub nadpisanej przez argumenty CLI")
numbered("agent/ppo_agent.py: train_ppo() - inicjalizacja srodowiska i modelu PPO")
numbered("environment/wrappers.py: make_sb3_env() - budowa srodowiska TrafficLightEnv i konwersja SuperSuit")
numbered("SB3 PPO.learn() - glowna petla treningowa")
numbered("Kazdy krok: traffic_env.step() wywoluje sample_arrivals(), resolve_phase_change(), process_departures()")
numbered("Co 1000 krokow: RedisMetricsCallback zapisuje metryki do Redis (jesli dostepny)")
numbered("Co 10 000 krokow: EvalCallback ocenia model na osobnym srodowisku ewaluacyjnym i zapisuje najlepszy checkpoint")
numbered("Po zakonczeniu: model.save() zapisuje finalny model do katalogu models/")

h2("2.3. Przeplyw danych podczas ewaluacji")
numbered("scripts/evaluate.py - wczytanie wytrenowanego modelu (agent/ppo_agent.py: load_model())")
numbered("Budowa srodowiska ewaluacyjnego (environment/wrappers.py: make_sb3_env())")
numbered("evaluate_policy() z biblioteki SB3 - uruchomienie N epizodow polityka PPO")
numbered("agent/baseline.py: evaluate_baseline() - uruchomienie tych samych epizodow polityka cykliczna (FixedCyclePolicy)")
numbered("Porownanie sredniej nagrody i odchylenia standardowego obu podejsc")

page_break()

# =====================================================================
# 3. STRUKTURA KATALOGOW PROJEKTU
# =====================================================================

h1("3. Struktura katalogow projektu")

para(
    "Ponizej przedstawiono pelna strukture katalogow i plikow repozytorium wraz z "
    "krotkim opisem przeznaczenia kazdego elementu. Szczegolowy opis zawartosci "
    "kazdego pliku znajduje sie w kolejnych rozdzialach."
)

code_block(
"""
agent-traffic-signal/
|-- config.py                  - centralna konfiguracja (Pydantic)
|-- main.py                     - plik startowy projektu (uv run)
|-- docker-compose.yml          - definicja uslugi Redis
|-- pyproject.toml               - metadane projektu i zaleznosci
|-- uv.lock                       - zablokowane wersje zaleznosci
|-- .python-version               - wymagana wersja Pythona (3.13)
|-- .gitignore                    - reguly ignorowania plikow w git
|-- README.md                     - dokumentacja uzytkownika (jezyk angielski)
|-- DEEP_DIVE.md                  - przewodnik techniczny dla zespolu
|-- PRESENTATION_PLAN.md          - plan prezentacji projektu
|
|-- environment/                  - warstwa symulacji i srodowiska RL
|   |-- __init__.py
|   |-- traffic_logic.py          - czyste funkcje symulacyjne
|   |-- traffic_env.py            - srodowisko PettingZoo (TrafficLightEnv)
|   |-- wrappers.py               - konwersja SuperSuit -> SB3 VecEnv
|
|-- agent/                        - warstwa agenta RL
|   |-- __init__.py
|   |-- ppo_agent.py              - trening PPO + callback Redis
|   |-- hyperopt.py               - optymalizacja hiperparametrow (Optuna)
|   |-- baseline.py               - polityka bazowa (fixed-cycle)
|
|-- state/                        - warstwa stanu
|   |-- __init__.py
|   |-- redis_manager.py          - operacje na Redis
|
|-- api/                          - warstwa API
|   |-- __init__.py
|   |-- server.py                 - serwer FastAPI
|
|-- visualization/                - warstwa wizualizacji
|   |-- __init__.py
|   |-- renderer.py               - widok skrzyzowania (matplotlib)
|   |-- plots.py                  - wykresy treningu i porownan
|
|-- scripts/                      - skrypty CLI
|   |-- train.py                  - trening agenta
|   |-- evaluate.py               - ewaluacja i porownanie z baseline
|   |-- run_api.py                - uruchomienie serwera FastAPI
|   |-- generate_plots.py         - generowanie wykresow do raportu
|
|-- tests/                        - testy automatyczne
|   |-- test_env.py               - testy srodowiska i logiki symulacji
|   |-- test_agent.py             - testy polityki bazowej
|   |-- test_api.py               - testy endpointow API
|
|-- notebooks/
|   |-- analysis.ipynb            - notatnik analizy wynikow
|
|-- models/                       - zapisane modele PPO (checkpointy)
|   |-- ppo_traffic_final.zip
|   |-- best/best_model.zip
|   |-- logs/evaluations.npz
|
|-- figures/                      - wygenerowane wykresy raportu
|
|-- resources/
|   |-- project-requirements.pdf  - wymagania projektu
|   |-- documentation.pdf         - dokumentacja zrodlowa
"""
)

h2("3.1. Przeglad katalogow")
make_table(
    ["Katalog", "Liczba plikow .py", "Odpowiedzialnosc"],
    [
        ["environment/", "3 + __init__", "Logika symulacji ruchu i sygnalizacji, definicja srodowiska PettingZoo, adaptacja do SB3"],
        ["agent/", "3 + __init__", "Trening agenta PPO, optymalizacja hiperparametrow, polityka bazowa"],
        ["state/", "1 + __init__", "Komunikacja ze sklepem stanu Redis (metryki, status, konfiguracja)"],
        ["api/", "1 + __init__", "REST API do zdalnego sterowania i monitorowania systemu"],
        ["visualization/", "2 + __init__", "Renderowanie graficzne skrzyzowania i wykresy analityczne"],
        ["scripts/", "4", "Punkty wejscia CLI - trening, ewaluacja, API, generowanie wykresow"],
        ["tests/", "3", "Testy jednostkowe i integracyjne (26 przypadkow testowych)"],
        ["notebooks/", "1 (ipynb)", "Koncowa analiza wynikow i wizualizacje do raportu"],
        ["models/", "-", "Persystencja wytrenowanych modeli i logow ewaluacji"],
        ["figures/", "-", "Wygenerowane wykresy uzywane w prezentacji i raporcie"],
        ["resources/", "-", "Materialy zrodlowe (wymagania projektu, dokumentacja)"],
    ],
    col_widths=[3.5, 3, 10],
)

page_break()

# =====================================================================
# 4. KONFIGURACJA - config.py
# =====================================================================

h1("4. Modul konfiguracyjny - config.py")

para(
    "Plik config.py stanowi centralny punkt konfiguracji calego systemu. Wykorzystuje "
    "biblioteke Pydantic do definicji modeli danych z walidacja typow oraz wartosciami "
    "domyslnymi. Dzieki temu konfiguracja moze byc latwo serializowana do formatu JSON "
    "(np. do przesylania przez API lub zapisu w Redis) oraz nadpisywana czesciowo "
    "(model_copy(update=...))."
)
para("Modul definiuje cztery modele konfiguracyjne: EnvConfig, TrainConfig, RedisConfig oraz AppConfig (agregujacy pozostale).")

h2("4.1. EnvConfig - konfiguracja srodowiska symulacyjnego")
make_table(
    ["Parametr", "Wartosc domyslna", "Opis"],
    [
        ["lambda_a", "0.6", "Parametr lambda procesu Poissona - intensywnosc naplywu pojazdow na drodze A (pojazdy/krok)"],
        ["lambda_b", "0.4", "Parametr lambda procesu Poissona dla drogi B"],
        ["green_drain_rate", "3", "Liczba pojazdow opuszczajacych kolejke w jednym kroku przy zielonym swietle"],
        ["min_green_steps", "4", "Minimalny czas trwania fazy zielonej (zapobiega migotaniu sygnalizacji)"],
        ["yellow_steps", "2", "Czas trwania fazy zoltej (przejsciowej) przy zmianie sygnalu"],
        ["max_steps", "500", "Liczba krokow w jednym epizodzie symulacji"],
        ["alpha", "0.5", "Wspolczynnik kooperacji w funkcji nagrody - waga kolejki drugiej drogi"],
        ["max_queue", "50", "Maksymalna dlugosc kolejki na jednej drodze (ograniczenie pojemnosci)"],
    ],
    col_widths=[4, 3, 10],
)

h2("4.2. TrainConfig - konfiguracja procesu treningu PPO")
make_table(
    ["Parametr", "Wartosc domyslna", "Opis"],
    [
        ["total_timesteps", "300 000", "Calkowita liczba krokow treningowych"],
        ["learning_rate", "3e-4", "Wspolczynnik uczenia optymalizatora"],
        ["n_steps", "2048", "Liczba krokow zbieranych przed kazda aktualizacja polityki (rollout)"],
        ["batch_size", "64", "Rozmiar minibatcha podczas gradientowej aktualizacji"],
        ["n_epochs", "10", "Liczba epok gradientowych na jednym rollout"],
        ["gamma", "0.99", "Wspolczynnik dyskontowania nagrod przyszlych"],
        ["gae_lambda", "0.95", "Parametr wygladzania GAE (Generalized Advantage Estimation)"],
        ["clip_range", "0.2", "Zakres przyciecia (clipping) w funkcji celu PPO"],
        ["ent_coef", "0.01", "Wspolczynnik bonusu entropii (eksploracja)"],
        ["vf_coef", "0.5", "Waga funkcji straty wartosci (value loss) w funkcji celu"],
        ["seed", "42", "Ziarno generatora liczb losowych"],
        ["eval_freq", "10 000", "Czestotliwosc ewaluacji modelu w trakcie treningu (w krokach)"],
        ["n_eval_episodes", "10", "Liczba epizodow uzywanych do ewaluacji"],
        ["model_dir", "\"models\"", "Katalog zapisu modeli i checkpointow"],
    ],
    col_widths=[4, 3, 10],
)

h2("4.3. RedisConfig - konfiguracja polaczenia z Redis")
make_table(
    ["Parametr", "Wartosc domyslna", "Opis"],
    [
        ["host", "\"localhost\"", "Adres serwera Redis"],
        ["port", "6379", "Port serwera Redis"],
        ["db", "0", "Numer bazy danych Redis"],
        ["state_key", "\"traffic:state\"", "Klucz przechowujacy aktualny stan srodowiska (hash)"],
        ["metrics_key", "\"traffic:metrics\"", "Klucz listy metryk treningu (kolejka czasowa)"],
        ["config_key", "\"traffic:config\"", "Klucz przechowujacy aktualna konfiguracje (JSON)"],
        ["status_key", "\"traffic:status\"", "Klucz statusu treningu (idle / starting / running / finished)"],
    ],
    col_widths=[4, 4.5, 8.5],
)

h2("4.4. AppConfig oraz DEFAULT_CONFIG")
para(
    "Model AppConfig agreguje wszystkie trzy konfiguracje czesciowe (env, train, redis) "
    "w jeden obiekt. Globalna instancja DEFAULT_CONFIG jest punktem wejscia uzywanym "
    "przez wszystkie skrypty CLI oraz serwer API. Wszystkie wartosci mozna nadpisywac "
    "lokalnie przy uzyciu metody model_copy(update={...}), co jest wykorzystywane np. "
    "w endpointcie PUT /config oraz w przeszukiwaniu hiperparametrow Optuna (parametr "
    "alpha jest tam nadpisywany dla kazdego trialu)."
)
code_block(
"""class AppConfig(BaseModel):
    env: EnvConfig = EnvConfig()
    train: TrainConfig = TrainConfig()
    redis: RedisConfig = RedisConfig()

DEFAULT_CONFIG = AppConfig()"""
)

page_break()

# =====================================================================
# 5. ENVIRONMENT/
# =====================================================================

h1("5. Warstwa symulacji - katalog environment/")

para(
    "Katalog environment/ zawiera implementacje srodowiska symulacyjnego skrzyzowania "
    "wraz z logika ruchu drogowego i sygnalizacji. Zostal podzielony na trzy pliki: "
    "czysta logike symulacji (traffic_logic.py), srodowisko zgodne ze standardem "
    "PettingZoo (traffic_env.py) oraz warstwe adaptacyjna do Stable-Baselines3 "
    "(wrappers.py)."
)

h2("5.1. environment/traffic_logic.py")
para(
    "Plik zawiera wylacznie czyste funkcje (bez zaleznosci od frameworkow RL), co "
    "ulatwia ich testowanie jednostkowe i analize matematyczna. Zaimplementowano trzy "
    "funkcje."
)

h3("5.1.1. sample_arrivals(lambda_a, lambda_b, rng)")
para(
    "Funkcja losuje liczbe pojazdow, ktore przybyly na obie drogi w danym kroku "
    "czasowym, korzystajac z rozkladu Poissona generatora numpy.random.Generator. "
    "Rozklad Poissona modeluje liczbe niezaleznych zdarzen (przyjazdow pojazdow) "
    "wystepujacych z pewna srednia intensywnoscia w jednostce czasu - jest to klasyczny "
    "przyklad procesu stochastycznego dyskretnego w czasie, bezposrednio zwiazany z "
    "tematyka przedmiotu Analiza regresji i szeregow czasowych."
)
code_block(
"""def sample_arrivals(lambda_a, lambda_b, rng):
    return int(rng.poisson(lambda_a)), int(rng.poisson(lambda_b))"""
)

h3("5.1.2. process_departures(queue, is_green, rate)")
para(
    "Funkcja modeluje odjazd pojazdow z kolejki. Jezeli sygnal jest zielony, z kolejki "
    "odejmowane jest do rate pojazdow (domyslnie 3), przy czym kolejka nigdy nie "
    "przyjmuje wartosci ujemnych. Jezeli sygnal jest czerwony lub zolty, kolejka nie "
    "zmienia sie."
)
code_block(
"""def process_departures(queue, is_green, rate):
    if not is_green:
        return queue
    return max(0, queue - rate)"""
)
para(
    "Wybor wartosci rate=3 nie jest przypadkowy - przy lambda_a=0.6 daje to mozliwosc "
    "szybkiego oprozniania kolejki w czasie trwania zielonego swiatla, co zapewnia "
    "istnienie stanu rownowagi (sredni naplyw < sredni odplyw przy zielonym)."
)

h3("5.1.3. resolve_phase_change(action_a, action_b, current_green, time_in_phase, min_green)")
para(
    "Jest to najistotniejsza funkcja calego srodowiska - implementuje reguly zmiany "
    "fazy sygnalizacji na podstawie akcji obu agentow. Zwraca pare (nastepny_zielony, "
    "czy_zmieniono)."
)
bullet("Minimalny czas zielonego: jezeli time_in_phase < min_green, zmiana jest zablokowana - zapobiega to migotaniu sygnalizacji.")
bullet("Akcja 1 oznacza żądanie zmiany fazy - moze je wystawic zarowno agent aktualnie zielony (chce ustąpić), jak i agent czerwony (zglasza potrzebe zielonego).")
bullet("Reguła anulowania (cancel-out): jezeli oba agenty zglosza żądanie zmiany jednoczesnie, żądania anuluja sie wzajemnie i faza nie jest zmieniana - mechanizm ten zapobiega oscylacjom.")
bullet("Jezeli wystapi rownoznacznie tylko jedno żądanie zmiany, sygnalizacja przelacza sie na druga droge.")
bullet("Jezeli zadna ze stron nie zglasza żądania, faza pozostaje bez zmian.")
code_block(
"""def resolve_phase_change(action_a, action_b, current_green, time_in_phase, min_green):
    if time_in_phase < min_green:
        return current_green, False

    wants_change_a = action_a == 1
    wants_change_b = action_b == 1

    if wants_change_a and wants_change_b:
        return current_green, False        # konflikt -> brak zmiany

    if current_green == "A" and (wants_change_a or wants_change_b):
        return "B", True
    if current_green == "B" and (wants_change_a or wants_change_b):
        return "A", True

    return current_green, False"""
)

h2("5.2. environment/traffic_env.py - klasa TrafficLightEnv")
para(
    "Glowna klasa srodowiska dziedziczy po pettingzoo.ParallelEnv i implementuje pelne "
    "API srodowiska wieloagentowego z dwoma agentami: light_A oraz light_B. Stan "
    "wewnetrzny srodowiska sklada sie z szesciu zmiennych:"
)
make_table(
    ["Zmienna", "Typ", "Opis"],
    [
        ["_queue_a", "int", "Liczba pojazdow czekajacych na drodze A"],
        ["_queue_b", "int", "Liczba pojazdow czekajacych na drodze B"],
        ["_current_green", "str", "Aktualna droga z zielonym swiatlem (\"A\" lub \"B\")"],
        ["_time_in_phase", "int", "Liczba krokow od ostatniej zmiany fazy"],
        ["_yellow_remaining", "int", "Licznik odliczajacy czas trwania zoltego swiatla"],
        ["_step_count", "int", "Numer aktualnego kroku w epizodzie"],
    ],
    col_widths=[3.5, 2, 11.5],
)

h3("5.2.1. Przestrzenie obserwacji i akcji")
para(
    "Funkcje observation_space() oraz action_space() sa dekorowane @functools.lru_cache, "
    "co jest wymagane przez API PettingZoo (przestrzenie musza byc niezmienne i "
    "hashowalne dla danego agenta). Kazdy agent otrzymuje 4-elementowy wektor "
    "obserwacji znormalizowany do przedzialu [0, 1]:"
)
code_block(
"""[own_queue / max_queue,        # zatloczenie wlasnej drogi (0-1)
 other_queue / max_queue,       # zatloczenie drugiej drogi (0-1)
 own_phase / 2.0,                # faza: 0=czerwone, 0.5=zielone, 1=zolte
 time_in_phase / max_steps]      # czas trwania aktualnej fazy (0-1)"""
)
para(
    "Akcja kazdego agenta jest dyskretna i binarna: 0 = pozostan w obecnej fazie, "
    "1 = zglos żądanie zmiany fazy."
)
para(
    "Normalizacja obserwacji jest kluczowa dla stabilnosci treningu sieci "
    "neuronowej - surowe wartosci kolejek (0-50) oraz czasu (0-500) zdominowalyby "
    "sygnal fazy (0-2), gdyby nie zostaly sprowadzone do wspolnej skali."
)
para(
    "Obserwacje sa symetryczne wzgledem agentow - kazdy z nich widzi swiat ze swojej "
    "perspektywy jako [wlasna_kolejka, druga_kolejka, ...]. Dzieki tej symetrii oba "
    "agenty moga korzystac z jednej, wspoldzielonej sieci neuronowej (parameter "
    "sharing)."
)

h3("5.2.2. Metoda reset()")
para(
    "Resetuje stan srodowiska: obie kolejki ustawiane sa na 0, aktywna droga "
    "ustawiana jest na A, liczniki czasu i fazy zoltej zerowane, a generator liczb "
    "losowych inicjalizowany jest opcjonalnym ziarnem (seed)."
)

h3("5.2.3. Metoda step(actions) - kolejnosc wykonywanych operacji")
numbered("Inkrementacja licznika kroku (_step_count)")
numbered("Losowanie przyjazdow Poissona i dodanie ich do obu kolejek (z ograniczeniem do max_queue)")
numbered("Jezeli aktywna jest faza zolta: dekrementacja licznika zoltego, pominiecie logiki zmiany fazy")
numbered("W przeciwnym razie: wywolanie resolve_phase_change() na podstawie akcji agentow")
numbered("Jezeli faza zostala zmieniona: ustawienie odliczania zoltego i resetowanie time_in_phase")
numbered("Jezeli faza nie zostala zmieniona: inkrementacja time_in_phase")
numbered("Przetworzenie odjazdow pojazdow (process_departures) - tylko poza faza zolta")
numbered("Obliczenie nagrod dla obu agentow (funkcja opisana w sekcji 5.2.4)")
numbered("Sprawdzenie warunku obciecia epizodu (truncation) po przekroczeniu max_steps")
numbered("Zwrocenie (obs, rewards, terminations, truncations, infos)")
para(
    "Brak odjazdow w fazie zoltej odwzorowuje rzeczywiste zachowanie ruchu drogowego "
    "(pojazdy zatrzymuja sie na zoltym swietle) i jednoczesnie wprowadza koszt "
    "kazdej zmiany sygnalizacji - agent traci 2 kroki przepustowosci przy kazdym "
    "przelaczeniu, co zapobiega nadmiernie czestym zmianom fazy."
)

h3("5.2.4. Funkcja nagrody")
para(
    "Nagroda kazdego agenta jest funkcja obu kolejek, znormalizowana do przedzialu "
    "[-1, 0], gdzie 0 oznacza stan idealny (obie kolejki puste), a -1 oznacza stan "
    "najgorszy (obie kolejki na poziomie max_queue)."
)
code_block(
"""scale = max_queue * (1.0 + alpha)               # = 50 * 1.5 = 75

reward_A = -(queue_A + alpha * queue_B) / scale
reward_B = -(queue_B + alpha * queue_A) / scale"""
)
para(
    "Wspolczynnik alpha (domyslnie 0.5) jest wspolczynnikiem kooperacji - okresla, "
    "w jakim stopniu kazdy agent powinien dbac o stan kolejki drugiej drogi. "
    "Przy alpha=0 agenty bylyby calkowicie egoistyczne i nigdy nie ustepowalyby "
    "zielonego swiatla; przy alpha=1 oba agenty traktowalyby obie kolejki rownorzednie."
)

h3("5.2.5. Cykl zycia epizodu i renderowanie")
para(
    "Epizod rozpoczyna sie od pustych kolejek i zielonego swiatla na drodze A, trwa "
    "max_steps=500 krokow i konczy sie poprzez obciecie (truncation), nie zaś "
    "osiagniecie stanu terminalnego - srodowisko nie definiuje warunku „wygranej”. "
    "Po obcieciu lista self.agents jest czyszczona, co jest sygnalem dla PettingZoo, "
    "ze epizod sie zakonczyl. Metoda render() obsluguje dwa rezimy: \"human\" "
    "(interaktywne okno matplotlib, TkAgg) oraz \"rgb_array\" (zwraca tablice numpy z "
    "obrazem klatki - wykorzystywane przy generowaniu zrzutow do raportu)."
)

h2("5.3. environment/wrappers.py - make_sb3_env()")
para(
    "Stable-Baselines3 oczekuje standardowego srodowiska Gymnasium VecEnv, natomiast "
    "PettingZoo udostepnia wieloagentowe ParallelEnv - te dwa API nie sa ze soba "
    "kompatybilne. Funkcja make_sb3_env() rozwiazuje ten problem przy uzyciu "
    "biblioteki SuperSuit w dwoch krokach:"
)
code_block(
"""env = TrafficLightEnv(env_config=env_config)
env.reset(seed=seed)

# Krok 1: splaszczenie srodowiska 2-agentowego do VecEnv z 2 pod-srodowiskami
env = ss.pettingzoo_env_to_vec_env_v1(env)

# Krok 2: agregacja do formatu kompatybilnego z SB3
env = ss.concat_vec_envs_v1(env, n_envs, num_cpus=1, base_class="stable_baselines3")"""
)
para(
    "Po tej konwersji SB3 widzi srodowisko VecEnv z num_envs = 2 (jedno na agenta). "
    "Kazde „pod-srodowisko” generuje wlasna obserwacje i akcje, ale wszystkie "
    "operuja na tej samej, wspolnej instancji TrafficLightEnv. W trakcie treningu SB3 "
    "podaje doswiadczenia obu agentow do jednej sieci neuronowej - jest to "
    "najprostsza forma uczenia wieloagentowego, znana jako parameter sharing "
    "(wspoldzielenie parametrow)."
)

h3("5.3.1. Lata bledu: brakujaca metoda seed()")
para(
    "Stable-Baselines3 wywoluje env.seed() podczas inicjalizacji, jednak klasa "
    "ConcatVecEnv z biblioteki SuperSuit nie implementuje tej metody, co powoduje "
    "blad AttributeError. Wrapper SB3 delegowal wywolanie seed() do atrybutu "
    "self.venv, dlatego latke nalezy nalozyc na wewnetrzny obiekt env.venv, nie na "
    "zewnetrzny wrapper SB3:"
)
code_block(
"""inner = env.venv if hasattr(env, "venv") else env
if not hasattr(inner, "seed"):
    inner.seed = lambda seed=None: [seed] * env.num_envs"""
)

page_break()

# =====================================================================
# 6. AGENT/
# =====================================================================

h1("6. Warstwa agenta - katalog agent/")

para(
    "Katalog agent/ zawiera trzy modulu odpowiedzialne za: trening agenta PPO "
    "(ppo_agent.py), optymalizacje hiperparametrow (hyperopt.py) oraz polityke "
    "bazowa do porownan (baseline.py)."
)

h2("6.1. agent/ppo_agent.py")

h3("6.1.1. Klasa RedisMetricsCallback")
para(
    "Callback dziedziczacy po stable_baselines3.common.callbacks.BaseCallback. Co "
    "1000 krokow (self.n_calls % 1000 == 0) odczytuje bufor ostatnich epizodow "
    "(model.ep_info_buffer), oblicza srednia nagrode i srednia dlugosc z 10 "
    "najnowszych epizodow i wysyla je do Redis przy pomocy "
    "redis_manager.push_metrics(). Dodatkowo, jezeli Redis jest dostepny, "
    "aktualizuje biezacy stan srodowiska (set_env_state)."
)
code_block(
"""class RedisMetricsCallback(BaseCallback):
    def _on_step(self) -> bool:
        if self.n_calls % 1000 == 0 and len(self.model.ep_info_buffer) > 0:
            recent = list(self.model.ep_info_buffer)[-10:]
            avg_reward = sum(ep["r"] for ep in recent) / len(recent)
            avg_length = sum(ep["l"] for ep in recent) / len(recent)
            self.redis_manager.push_metrics({
                "timestep": self.num_timesteps,
                "avg_reward": round(avg_reward, 2),
                "avg_length": round(avg_length, 2),
            })
        return True"""
)

h3("6.1.2. Funkcja train_ppo(config, redis_manager)")
para("Glowna funkcja treningowa wykonuje nastepujace kroki:")
numbered("Tworzy katalog models/ (jesli nie istnieje)")
numbered("Buduje srodowisko treningowe i ewaluacyjne (rozne ziarna: seed oraz seed+1000)")
numbered("Inicjalizuje model PPO z polityka MlpPolicy oraz wszystkimi hiperparametrami z TrainConfig")
numbered("Dodaje EvalCallback - ewaluacja co eval_freq krokow, zapis najlepszego modelu do models/best/")
numbered("Opcjonalnie dodaje RedisMetricsCallback, jesli przekazano redis_manager")
numbered("Ustawia status treningu w Redis na \"running\" przed startem i \"finished\" po zakonczeniu")
numbered("Wywoluje model.learn(total_timesteps=..., callback=callbacks)")
numbered("Zapisuje finalny model do models/ppo_traffic_final")

h3("6.1.3. Hiperparametry PPO uzyte w projekcie")
make_table(
    ["Hiperparametr", "Wartosc", "Znaczenie"],
    [
        ["learning_rate", "3e-4", "Tempo aktualizacji wag sieci - kompromis miedzy stabilnoscia a szybkoscia uczenia"],
        ["n_steps", "2048", "Liczba krokow zbieranych przed kazda aktualizacja (rollout)"],
        ["batch_size", "64", "Rozmiar minibatcha w gradientowej aktualizacji"],
        ["n_epochs", "10", "Liczba przejsc gradientowych po jednym rollout"],
        ["gamma", "0.99", "Wspolczynnik dyskontowania - horyzont ok. 100 krokow w przod"],
        ["gae_lambda", "0.95", "Wygladzanie GAE - balans miedzy obciazeniem a wariancja estymacji"],
        ["clip_range", "0.2", "PPO clipping - ogranicza zmiane polityki do +/-20% na aktualizacje"],
        ["ent_coef", "0.01", "Bonus entropii - zacheca do eksploracji"],
        ["vf_coef", "0.5", "Waga funkcji straty wartosci w funkcji celu"],
    ],
    col_widths=[3.5, 2.5, 11],
)

h3("6.1.4. Funkcja load_model(path)")
para("Prosta funkcja pomocnicza opakowujaca PPO.load(path) z biblioteki Stable-Baselines3, uzywana w skryptach evaluate.py, generate_plots.py oraz w API.")

h2("6.2. agent/hyperopt.py - optymalizacja hiperparametrow")
para(
    "Modul integruje biblioteke Optuna w celu automatycznego przeszukiwania "
    "przestrzeni hiperparametrow PPO przy uzyciu algorytmu Tree-structured Parzen "
    "Estimator (TPE) - probabilistycznego modelu Bayesowskiego, ktory na podstawie "
    "wynikow poprzednich prob (trials) decyduje, w ktorym obszarze przestrzeni "
    "parametrow nalezy probkowac kolejne kombinacje."
)

h3("6.2.1. Funkcja objective(trial, base_config)")
para("Dla kazdego triala Optuna proponuje wartosci hiperparametrow, trenuje model PPO przez 100 000 krokow i zwraca srednia nagrode z ewaluacji jako wartosc do maksymalizacji.")
make_table(
    ["Hiperparametr", "Przestrzen przeszukiwania", "Typ rozkladu"],
    [
        ["learning_rate", "1e-5 - 1e-3", "logarytmiczny"],
        ["n_steps", "{256, 512, 1024, 2048}", "kategoryczny"],
        ["batch_size", "{32, 64, 128, 256}", "kategoryczny"],
        ["n_epochs", "3 - 15", "calkowitoliczbowy"],
        ["gamma", "0.95 - 0.999", "ciagly"],
        ["gae_lambda", "0.8 - 0.99", "ciagly"],
        ["clip_range", "0.1 - 0.4", "ciagly"],
        ["ent_coef", "1e-4 - 0.1", "logarytmiczny"],
        ["alpha", "0.0 - 1.0", "ciagly (wspolczynnik kooperacji)"],
    ],
    col_widths=[3.5, 5, 8.5],
)
para(
    "Wlaczenie wspolczynnika alpha do przestrzeni przeszukiwania jest istotnym "
    "elementem badawczym projektu - pozwala odpowiedziec na pytanie, jaki poziom "
    "kooperacji miedzy agentami jest optymalny dla minimalizacji laczne kolejki."
)

h3("6.2.2. Funkcja run_hyperopt(n_trials, storage, base_config)")
para(
    "Tworzy (lub wczytuje, jezeli istnieje) study Optuna o nazwie traffic_ppo, "
    "z kierunkiem optymalizacji maximize, przechowywane w bazie SQLite "
    "(sqlite:///optuna_study.db). Nastepnie wywoluje study.optimize() z funkcja "
    "objective dla zadanej liczby trialow."
)
para(
    "Liczba krokow na trial (100 000) zostala wybrana jako kompromis - pelny trening "
    "trwa 300 000 krokow (ok. 65 sekund na CPU), a przy 50 trialach po 300 000 "
    "krokow optymalizacja zajelaby ok. 54 minut. Przy 100 000 krokow na trial czas "
    "skraca sie do ok. 18 minut, przy zachowaniu wzgledem stabilnego rankingu "
    "hiperparametrow."
)

h2("6.3. agent/baseline.py - polityka bazowa (fixed-cycle)")

h3("6.3.1. Klasa FixedCyclePolicy")
para(
    "Implementuje prosty, sztywny harmonogram cykliczny - przelaczenie fazy nastepuje "
    "co cycle_length krokow (domyslnie 15), ale wylacznie jezeli agent jest aktualnie "
    "zielony. Agent czerwony zawsze zwraca akcje 0 (pozostan), co zapobiega "
    "wlaczeniu reguly anulowania (cancel-out) opisanej w sekcji 5.1.3."
)
code_block(
"""class FixedCyclePolicy:
    def predict(self, obs):
        own_phase = obs[2]
        is_green = abs(own_phase - 0.5) < 0.01
        self._step += 1
        if is_green and self._step % self.cycle_length == 0:
            return 1
        return 0"""
)
para(
    "Krytyczna decyzja projektowa: gdyby oba agenty (A i B) korzystaly z identycznego "
    "licznika cyklu i obie zglaszaly żądanie zmiany w tym samym kroku, reguła "
    "anulowania zablokowalaby kazda zmiane fazy - baseline nigdy by sie nie "
    "przelaczal. Dlatego tylko agent zielony jest aktywny."
)

h3("6.3.2. Funkcja evaluate_baseline(cycle_length, n_episodes, env_config, seed)")
para(
    "Uruchamia n_episodes epizodow srodowiska TrafficLightEnv (bezposrednio, bez "
    "wektoryzacji SB3), kontrolujac obie sygnalizacje przy uzyciu dwoch instancji "
    "FixedCyclePolicy. Zbiera sume nagrod oraz srednia dlugosc kolejki w kazdym "
    "epizodzie i zwraca slownik z metrykami: mean_reward, std_reward, mean_queue, "
    "std_queue, n_episodes, cycle_length."
)
para(
    "Cykl o dlugosci 15 krokow zapewnia kazdej drodze sprawiedliwy udzial czasu "
    "zielonego, jednak harmonogram ten nie reaguje na rzeczywiste natezenie ruchu - "
    "co stanowi punkt odniesienia, wzgledem ktorego oceniana jest skutecznosc "
    "adaptacyjnego agenta PPO."
)

page_break()

# =====================================================================
# 7. STATE/ I API/
# =====================================================================

h1("7. Warstwa stanu i API")

h2("7.1. state/redis_manager.py - klasa RedisStateManager")
para(
    "Klasa RedisStateManager opakowuje klienta redis.Redis i udostepnia wyzszej "
    "poziomu interfejs do odczytu i zapisu stanu systemu. Polaczenie z Redis jest "
    "tworzone leniwie (lazy initialization) przy pierwszym dostepie do wlasciwosci "
    "client, z opcja decode_responses=True (klucze i wartosci odczytywane jako "
    "lancuchy znakow, nie bajty)."
)
para("Redis przechowuje cztery rodzaje danych:")
make_table(
    ["Klucz Redis", "Typ struktury", "Zawartosc"],
    [
        ["traffic:state", "Hash", "Biezacy stan srodowiska (kolejki, faza, krok)"],
        ["traffic:metrics", "List", "Szereg czasowy metryk treningu (timestep, avg_reward, avg_length)"],
        ["traffic:config", "String (JSON)", "Biezaca konfiguracja systemu"],
        ["traffic:status", "String", "Status treningu: idle / starting / running / finished"],
    ],
    col_widths=[4, 3.5, 8.5],
)
para("Wykorzystanie Redis pozwala oddzielic proces treningu od procesu API monitorujacego - proces treningowy zapisuje metryki, a API odczytuje je niezaleznie, bez koniecznosci dzielenia pamieci procesu.")

h3("7.1.1. Metody klasy RedisStateManager")
make_table(
    ["Metoda", "Opis"],
    [
        ["is_connected()", "Sprawdza dostepnosc Redis poprzez PING; przechwytuje redis.ConnectionError"],
        ["set_env_state(state)", "Zapisuje stan srodowiska jako hash (HSET); wartosci zlozone serializowane do JSON"],
        ["get_env_state()", "Odczytuje caly hash stanu (HGETALL)"],
        ["push_metrics(metrics)", "Dodaje wpis metryk na koniec listy (RPUSH, JSON)"],
        ["get_metrics(start, end)", "Odczytuje zakres listy metryk (LRANGE) i deserializuje JSON"],
        ["clear_metrics()", "Usuwa liste metryk (DEL) - wywolywane przed nowym treningiem"],
        ["set_config(config)", "Zapisuje konfiguracje jako JSON (SET)"],
        ["get_config()", "Odczytuje i deserializuje konfiguracje"],
        ["set_status(status)", "Zapisuje status treningu (SET)"],
        ["get_status()", "Odczytuje status treningu, domyslnie \"idle\""],
    ],
    col_widths=[5, 11],
)

h2("7.2. api/server.py - serwer FastAPI")
para(
    "Modul definiuje REST API o nazwie Traffic Light RL API (wersja 1.0.0). Serwer "
    "tworzy globalna, glebokow kopiowana instancje konfiguracji (_config), klienta "
    "Redis (_redis) oraz referencje do wątku treningowego (_train_thread)."
)

h3("7.2.1. Lista endpointow")
make_table(
    ["Metoda", "Sciezka", "Opis"],
    [
        ["GET", "/state", "Zwraca biezacy stan srodowiska z Redis (kolejki, faza); blad 503 jezeli Redis niedostepny"],
        ["GET", "/metrics", "Historia metryk treningu (parametry start, end pozwalaja na zakres)"],
        ["POST", "/train/start", "Uruchamia trening PPO w tle (osobny wątek); blad 409 jezeli trening juz trwa"],
        ["GET", "/train/status", "Zwraca status treningu (idle/starting/running/finished)"],
        ["GET", "/config", "Zwraca biezaca konfiguracje (env, train, redis)"],
        ["PUT", "/config", "Aktualizuje konfiguracje (czesciowo lub w calosci) i zapisuje w Redis"],
        ["POST", "/evaluate", "Uruchamia ewaluacje wytrenowanego modelu, opcjonalnie z porownaniem do baseline"],
        ["GET", "/hyperopt/results", "Zwraca najlepsze parametry i wartosc z badania Optuna (404 jezeli brak wynikow)"],
    ],
    col_widths=[2, 3.5, 10.5],
)

h3("7.2.2. Kluczowe decyzje projektowe")
bullet("Trening w osobnym wątku demona (threading.Thread, daemon=True) - endpoint POST /train/start natychmiast zwraca odpowiedz, a API pozostaje responsywne w trakcie wielogodzinnego treningu.")
bullet("Leniwe importy (lazy imports) - kosztowne moduly (torch, Stable-Baselines3) sa importowane wewnatrz funkcji endpointow, nie na poziomie modulu - dzieki temu start serwera jest natychmiastowy.")
bullet("Łagodna degradacja (graceful degradation) - w przypadku braku polaczenia z Redis wiekszosc endpointow zwraca sensowne wartosci domyslne lub blad HTTP 503/404 zamiast nieobslugiwanego wyjatku.")

h3("7.2.3. Model ConfigUpdate")
para(
    "Endpoint PUT /config przyjmuje model ConfigUpdate (Pydantic) z opcjonalnymi "
    "polami env: EnvConfig | None i train: TrainConfig | None. Aktualizacja jest "
    "stosowana czesciowo - tylko przekazane pola sa nadpisywane (config.model_copy"
    "(update=...)), pozostale zachowuja poprzednie wartosci."
)

h3("7.2.4. Endpoint POST /evaluate - szczegoly implementacji")
para(
    "Endpoint sprawdza istnienie pliku models/ppo_traffic_final.zip (blad 404 jezeli "
    "nie istnieje), wczytuje model, buduje srodowisko ewaluacyjne z ziarnem 99, "
    "wywoluje evaluate_policy() z SB3 i (opcjonalnie) evaluate_baseline(). Wynik "
    "zawiera zaokrąglone do dwóch miejsc po przecinku wartosci mean_reward i "
    "std_reward dla PPO oraz, jesli wybrano, analogiczne metryki dla baseline."
)

page_break()

# =====================================================================
# 8. VISUALIZATION/
# =====================================================================

h1("8. Warstwa wizualizacji - katalog visualization/")

h2("8.1. visualization/renderer.py")
para(
    "Plik zawiera funkcje renderujace widok skrzyzowania „z lotu ptaka” przy uzyciu "
    "matplotlib.patches. Funkcja pomocnicza _draw_intersection(ax, queue_a, queue_b, "
    "green, yellow, step) rysuje na przekazanych osiach: dwie przecinajace sie drogi, "
    "pole skrzyzowania, oznaczenia pasow, dwa sygnalizatory swietlne (kolka) oraz "
    "paski reprezentujace dlugosci kolejek na obu drogach. Kolor sygnalizatorow "
    "zalezy od stanu: zielony/czerwony w zaleznosci od aktywnej drogi, zolty w "
    "trakcie fazy przejsciowej."
)
make_table(
    ["Funkcja", "Tryb", "Opis"],
    [
        ["render_intersection(...)", "human", "Renderuje interaktywne okno matplotlib (backend TkAgg), aktualizowane na zywo (plt.pause)"],
        ["render_intersection_rgb(...)", "rgb_array", "Renderuje pojedyncza klatke do tablicy numpy (RGB), wykorzystywana do generowania zrzutow statycznych"],
    ],
    col_widths=[5.5, 2.5, 8],
)
para(
    "Funkcje renderujace sa wywolywane przez TrafficLightEnv.render() w zaleznosci "
    "od parametru render_mode przekazanego przy tworzeniu srodowiska, a takze "
    "bezposrednio przez scripts/generate_plots.py do generowania zrzutow skrzyzowania "
    "w roznych momentach epizodu."
)

h2("8.2. visualization/plots.py")
para("Modul zawiera trzy funkcje generujace wykresy analityczne na podstawie zebranych danych.")

h3("8.2.1. plot_training_curves(metrics, save_path, show)")
para(
    "Generuje wykres skladajacy sie z dwoch paneli: srednia nagroda w funkcji liczby "
    "krokow treningowych (timestep) oraz srednia dlugosc epizodu. Dane pochodza z "
    "listy metryk zapisanych w Redis przez RedisMetricsCallback. Jezeli lista metryk "
    "jest pusta, funkcja wypisuje ostrzezenie i nie generuje wykresu."
)

h3("8.2.2. plot_comparison(ppo_rewards, baseline_rewards, save_path, show)")
para(
    "Generuje wykres slupkowy porownujacy srednia nagrode (z odchyleniem "
    "standardowym jako wąsy bledu) miedzy agentem PPO a baseline cyklicznym. Na "
    "kazdym slupku wypisywana jest wartosc liczbowa sredniej."
)

h3("8.2.3. plot_optuna_results(study, save_path, show)")
para(
    "Wykorzystuje moduly optuna.visualization.matplotlib do wygenerowania dwoch "
    "wykresow: historii optymalizacji (plot_optimization_history) oraz istotnosci "
    "poszczegolnych hiperparametrow (plot_param_importances). Jezeli modul "
    "wizualizacji Optuna nie jest dostepny, funkcja wypisuje stosowny komunikat "
    "i nie przerywa wykonania."
)

page_break()

# =====================================================================
# 9. SCRIPTS/
# =====================================================================

h1("9. Skrypty uruchomieniowe - katalog scripts/")

para(
    "Katalog scripts/ zawiera cztery punkty wejscia wiersza polecen (CLI). Kazdy "
    "skrypt na poczatku dodaje katalog glowny projektu do sys.path, co pozwala na "
    "import modulow (config, environment, agent, ...) niezaleznie od miejsca "
    "wywolania skryptu."
)

h2("9.1. scripts/train.py")
para("Skrypt treningowy przyjmuje nastepujace argumenty wiersza polecen:")
make_table(
    ["Argument", "Domyslnie", "Opis"],
    [
        ["--timesteps", "None (300 000)", "Calkowita liczba krokow treningowych"],
        ["--seed", "None (42)", "Ziarno generatora liczb losowych"],
        ["--hyperopt", "False", "Uruchamia optymalizacje Optuna zamiast pojedynczego treningu"],
        ["--hyperopt-trials", "50", "Liczba prob (trials) przy optymalizacji hiperparametrow"],
        ["--no-redis", "False", "Wylacza logowanie metryk do Redis"],
    ],
    col_widths=[3.5, 3.5, 9],
)
para(
    "Jezeli flaga --no-redis nie jest podana, skrypt probuje polaczyc sie z Redis "
    "(RedisStateManager.is_connected()); w przypadku niepowodzenia trening "
    "kontynuowany jest bez logowania metryk, z odpowiednim ostrzezeniem na konsoli. "
    "Po zakonczeniu treningu model jest zapisywany do katalogu wskazanego w "
    "TrainConfig.model_dir (domyslnie models/)."
)
code_block(
"""# Domyslny trening (300k krokow)
uv run python scripts/train.py

# Trening niestandardowej dlugosci, bez Redis
uv run python scripts/train.py --timesteps 50000 --no-redis

# Optymalizacja hiperparametrow Optuna
uv run python scripts/train.py --hyperopt --hyperopt-trials 50"""
)

h2("9.2. scripts/evaluate.py")
para("Skrypt ewaluacyjny posiada nastepujace argumenty:")
make_table(
    ["Argument", "Domyslnie", "Opis"],
    [
        ["--model", "models/ppo_traffic_final", "Sciezka do wytrenowanego modelu"],
        ["--episodes", "20", "Liczba epizodow ewaluacyjnych"],
        ["--compare-baseline", "False", "Porownanie z polityka cykliczna"],
        ["--render", "False", "Renderowanie epizodow w trybie interaktywnym"],
        ["--seed", "42", "Ziarno generatora liczb losowych"],
    ],
    col_widths=[3.5, 5, 7.5],
)
para(
    "Funkcja _evaluate_headless() korzysta z evaluate_policy() (SB3) na "
    "zwektoryzowanym srodowisku i wypisuje srednia nagrode +/- odchylenie "
    "standardowe; opcjonalnie wywoluje evaluate_baseline() i wypisuje roznice "
    "(\"PPO improvement over baseline\")."
)
para(
    "Funkcja _evaluate_with_render() dziala bezposrednio na TrafficLightEnv "
    "(render_mode=\"human\") i dla kazdego agenta wykonuje niezalezna predykcje:"
)
code_block(
"""action_a, _ = model.predict(obs["light_A"], deterministic=True)
action_b, _ = model.predict(obs["light_B"], deterministic=True)
actions = {"light_A": int(action_a), "light_B": int(action_b)}"""
)
para(
    "To rozroznienie jest istotne - wczesniejsza, blednna implementacja wywolywala "
    "predykcje tylko jeden raz i stosowala ten sam wynik dla obu agentow, co "
    "prowadzilo do uruchomienia reguly anulowania (gdy obaj zwracali akcje 1) i "
    "calkowitego braku zmian fazy (opisane szerzej w rozdziale 15)."
)

h2("9.3. scripts/run_api.py")
para(
    "Minimalny skrypt startujacy serwer FastAPI przy uzyciu Uvicorn: "
    "uvicorn.run(\"api.server:app\", host=\"0.0.0.0\", port=8000, reload=True). "
    "Po starcie dokumentacja interaktywna API (Swagger UI) jest dostepna pod "
    "adresem http://localhost:8000/docs."
)
code_block("uv run python scripts/run_api.py")

h2("9.4. scripts/generate_plots.py")
para(
    "Skrypt generujacy wszystkie wykresy wykorzystywane w raporcie i prezentacji. "
    "Wczytuje wytrenowany model PPO oraz tworzy katalog figures/ (jesli nie "
    "istnieje). Wykorzystuje backend matplotlib \"Agg\" (bez okna graficznego), co "
    "pozwala uruchamiac skrypt w trybie headless (np. na serwerze CI)."
)
para("Skrypt generuje cztery pliki graficzne:")
make_table(
    ["Plik wynikowy", "Zawartosc"],
    [
        ["figures/queue_dynamics_comparison.png", "4 panele: dynamika kolejek A/B oraz harmonogram faz dla PPO i baseline (jeden epizod, seed=42)"],
        ["figures/reward_comparison.png", "Wykres slupkowy: srednia suma nagrod PPO vs baseline (20 epizodow, seed 100-119)"],
        ["figures/cumulative_reward.png", "Skumulowana nagroda w funkcji kroku epizodu - PPO vs baseline (seed=42)"],
        ["figures/intersection_snapshots.png", "Trzy zrzuty wizualizacji skrzyzowania (kroki 1, 100, 400) renderowane przy uzyciu render_intersection_rgb()"],
    ],
    col_widths=[6, 10],
)
para(
    "Funkcja run_episode(policy_fn, seed) jest funkcja pomocnicza wspolna dla "
    "wszystkich wykresow - uruchamia jeden epizod TrafficLightEnv przy przekazanej "
    "funkcji polityki (ppo_policy lub baseline_policy) i zbiera szeregi czasowe "
    "dlugosci kolejek, nagrod i aktywnej fazy."
)

page_break()

# =====================================================================
# 10. TESTS/
# =====================================================================

h1("10. Testy automatyczne - katalog tests/")

para(
    "Projekt zawiera 26 przypadkow testowych podzielonych na trzy pliki, "
    "uruchamiane przy pomocy pytest. Testy obejmuja logike symulacji, zgodnosc ze "
    "standardem PettingZoo, polityke bazowa oraz endpointy API."
)

h2("10.1. tests/test_env.py")
make_table(
    ["Klasa testowa", "Test", "Co weryfikuje"],
    [
        ["TestSampleArrivals", "test_returns_non_negative", "Liczba przyjazdow jest zawsze >= 0 dla 100 prob"],
        ["TestSampleArrivals", "test_mean_approximates_lambda", "Srednia z 10 000 prob jest bliska wartosciom lambda (E[X]=lambda)"],
        ["TestProcessDepartures", "test_green_drains", "Przy zielonym swietle kolejka zmniejsza sie o rate"],
        ["TestProcessDepartures", "test_red_no_drain", "Przy czerwonym swietle kolejka nie zmienia sie"],
        ["TestProcessDepartures", "test_no_negative_queue", "Kolejka nie spada poniżej zera"],
        ["TestProcessDepartures", "test_empty_queue", "Pusta kolejka pozostaje pusta"],
        ["TestResolvePhaseChange", "test_min_green_prevents_change", "Zmiana fazy zablokowana przed uplywem min_green"],
        ["TestResolvePhaseChange", "test_change_from_a_to_b", "Poprawna zmiana fazy z A na B"],
        ["TestResolvePhaseChange", "test_change_from_b_to_a", "Poprawna zmiana fazy z B na A"],
        ["TestResolvePhaseChange", "test_no_change_when_action_keep", "Brak zmiany, gdy obie akcje = 0"],
        ["TestResolvePhaseChange", "test_red_agent_can_request_green", "Agent czerwony moze wymusic zmiane fazy"],
        ["TestResolvePhaseChange", "test_both_request_cancels_out", "Reguła anulowania przy konflikcie żądań"],
        ["TestTrafficLightEnv", "test_reset", "reset() zwraca obserwacje dla obu agentow o ksztalcie (4,)"],
        ["TestTrafficLightEnv", "test_step_returns_correct_structure", "step() zwraca prawidlowa strukture (obs, rewards, ...)"],
        ["TestTrafficLightEnv", "test_episode_truncates", "Epizod konczy sie (self.agents == []) po max_steps krokach"],
        ["TestTrafficLightEnv", "test_observation_space_compliance", "Obserwacje naleza do zadeklarowanej przestrzeni obserwacji"],
        ["TestTrafficLightEnv", "test_rewards_are_negative", "Nagrody sa <= 0 (kara za kolejki)"],
        ["TestPettingZooCompliance", "test_parallel_api", "Oficjalny test zgodnosci API PettingZoo (parallel_api_test, 50 cykli)"],
    ],
    col_widths=[4.5, 5.5, 6],
)

h2("10.2. tests/test_agent.py")
make_table(
    ["Klasa testowa", "Test", "Co weryfikuje"],
    [
        ["TestFixedCyclePolicy", "test_returns_valid_actions", "predict() zawsze zwraca akcje 0 lub 1"],
        ["TestFixedCyclePolicy", "test_changes_at_cycle_boundary_when_green", "Zmiana zglaszana co cycle_length krokow przy zielonym"],
        ["TestFixedCyclePolicy", "test_no_change_when_red", "Agent czerwony nigdy nie zglasza zmiany"],
        ["TestFixedCyclePolicy", "test_reset", "reset() prawidlowo zeruje licznik kroku"],
        ["TestEvaluateBaseline", "test_returns_metrics", "evaluate_baseline() zwraca slownik z wymaganymi kluczami metryk"],
    ],
    col_widths=[4.5, 7, 4.5],
)
para(
    "Testy korzystaja z funkcji pomocniczych _green_obs() oraz _red_obs(), ktore "
    "tworza syntetyczne wektory obserwacji odpowiadajace agentowi w fazie zielonej "
    "(obs[2]=0.5) lub czerwonej (obs[2]=0.0)."
)

h2("10.3. tests/test_api.py")
make_table(
    ["Klasa testowa", "Test", "Co weryfikuje"],
    [
        ["TestConfigEndpoints", "test_get_config", "GET /config zwraca strukture z polami env, train, redis"],
        ["TestConfigEndpoints", "test_update_config", "PUT /config poprawnie aktualizuje wartosci lambda_a / lambda_b"],
        ["TestTrainStatus", "test_train_status", "GET /train/status zwraca pole status"],
    ],
    col_widths=[4.5, 6, 5.5],
)
para(
    "Testy API korzystaja z fastapi.testclient.TestClient, ktory uruchamia "
    "aplikacje FastAPI w pamieci, bez potrzeby startowania prawdziwego serwera "
    "HTTP."
)

h2("10.4. Wynik przebiegu testow")
para(
    "Pelny zestaw 26 testow zostal uruchomiony lokalnie poleceniem "
    "uv run python -m pytest tests/ -v. Wszystkie testy zakonczyly sie powodzeniem."
)
make_table(
    ["Plik testowy", "Liczba testow", "Status"],
    [
        ["tests/test_env.py", "18", "PASSED"],
        ["tests/test_agent.py", "5", "PASSED"],
        ["tests/test_api.py", "3", "PASSED"],
        ["SUMA", "26", "26 passed"],
    ],
    col_widths=[6, 5, 5],
)

page_break()

# =====================================================================
# 11. NOTEBOOKS/
# =====================================================================

h1("11. Notatnik analityczny - notebooks/analysis.ipynb")

para(
    "Notatnik Jupyter analysis.ipynb stanowi koncowy etap pipeline'u - laczy wyniki "
    "treningu, ewaluacji oraz optymalizacji hiperparametrow w jeden dokument "
    "analityczny, wykorzystywany do przygotowania raportu i prezentacji. Notatnik "
    "dodaje katalog glowny projektu do sys.path (sys.path.insert(0, '..')), aby "
    "mial dostep do wszystkich modulow."
)

h2("11.1. Sekcja 1 - Krzywe treningowe")
para(
    "Wykorzystuje RedisStateManager.get_metrics() do odczytania historii metryk "
    "zapisanych podczas treningu, a nastepnie plot_training_curves() do "
    "wygenerowania wykresu sredniej nagrody i dlugosci epizodu w funkcji liczby "
    "krokow (zapis do figures/training_curves.png)."
)

h2("11.2. Sekcja 2 - Agent RL vs baseline cykliczny")
para(
    "Wczytuje wytrenowany model (load_model('models/ppo_traffic_final')), tworzy "
    "srodowisko ewaluacyjne (make_sb3_env(seed=99)) i przez 20 epizodow zbiera "
    "skumulowana nagrode agenta PPO. Rownolegle wywoluje evaluate_baseline"
    "(n_episodes=20) dla polityki cyklicznej. Wyniki obu podejsc sa zestawiane na "
    "wykresie slupkowym plot_comparison() (zapis do figures/comparison.png)."
)

h2("11.3. Sekcja 3 - Wyniki optymalizacji hiperparametrow")
para(
    "Wczytuje (jesli istnieje) badanie Optuna optuna.load_study(study_name="
    "'traffic_ppo', storage='sqlite:///../optuna_study.db') i wypisuje najlepsze "
    "parametry oraz najlepsza wartosc funkcji celu. Jezeli plik bazy danych nie "
    "istnieje (np. optymalizacja nie zostala uruchomiona), notatnik wypisuje "
    "komunikat informacyjny i nie przerywa wykonania - blok jest zabezpieczony "
    "konstrukcja try/except."
)

h2("11.4. Sekcja 4 - Analiza dynamiki srodowiska")
para(
    "Tworzy nowa instancje TrafficLightEnv, resetuje ja z ziarnem 42, a nastepnie "
    "w petli while env.agents: wykonuje krok srodowiska przy uzyciu predykcji "
    "wytrenowanego modelu (z tym samym wektorem akcji dla obu agentow - jest to "
    "uproszczona wersja analizy, oparta na pojedynczej predykcji). Zbierane sa "
    "szeregi czasowe dlugosci kolejek obu drog oraz nagrody agenta light_A. Wynik "
    "przedstawiany jest na dwoch panelach (dynamika kolejek i nagroda w czasie) i "
    "zapisywany do figures/queue_dynamics.png."
)

page_break()

# =====================================================================
# 12. PLIKI POMOCNICZE I KONFIGURACYJNE
# =====================================================================

h1("12. Pliki pomocnicze i konfiguracyjne")

h2("12.1. main.py")
para(
    "Plik wejsciowy projektu, generowany przez szablon uv init. Zawiera funkcje "
    "main(), ktora wypisuje powitanie (\"Hello from agent-traffic-signal!\"). "
    "Faktycznym punktem wejscia do funkcjonalnosci projektu sa skrypty w katalogu "
    "scripts/ (train.py, evaluate.py, run_api.py, generate_plots.py)."
)

h2("12.2. pyproject.toml")
para(
    "Definiuje metadane pakietu (nazwa: agent-traffic-signal, wersja: 0.1.0) oraz "
    "minimalna wymagana wersje Pythona (>=3.13). Zawiera pelna liste zaleznosci "
    "produkcyjnych projektu:"
)
make_table(
    ["Pakiet", "Minimalna wersja", "Zastosowanie"],
    [
        ["fastapi", ">=0.135.1", "REST API"],
        ["gymnasium", ">=1.2.3", "Bazowe API przestrzeni obserwacji/akcji"],
        ["httpx", ">=0.28.1", "Klient HTTP (testy API)"],
        ["matplotlib", ">=3.10.8", "Wizualizacje"],
        ["numpy", ">=2.4.3", "Operacje numeryczne, rozklad Poissona"],
        ["optuna", ">=4.8.0", "Optymalizacja hiperparametrow"],
        ["pettingzoo", ">=1.25.0", "Srodowisko wieloagentowe"],
        ["pydantic", ">=2.12.5", "Modele konfiguracji i walidacja danych"],
        ["pytest", ">=9.0.2", "Testy automatyczne"],
        ["redis", ">=7.3.0", "Klient bazy Redis"],
        ["stable-baselines3", ">=2.7.1", "Implementacja PPO"],
        ["supersuit", ">=3.10.0", "Konwersja PettingZoo -> SB3 VecEnv"],
        ["torch", ">=2.10.0", "Backend obliczeniowy sieci neuronowych (PPO)"],
        ["uvicorn", ">=0.42.0", "Serwer ASGI dla FastAPI"],
    ],
    col_widths=[4, 3.5, 8.5],
)

h2("12.3. docker-compose.yml")
para(
    "Definiuje pojedyncza usluge redis (obraz redis:7-alpine), wystawiona na "
    "standardowym porcie 6379, z wolumenem redis_data zapewniajacym trwalosc danych "
    "miedzy restartami kontenera. Usluga jest opcjonalna - system dziala rowniez "
    "bez Redis (z ograniczona funkcjonalnoscia: brak metryk live, endpointy "
    "zwracajace status 503)."
)
code_block(
"""services:
  redis:
    image: redis:7-alpine
    ports:
      - "6379:6379"
    volumes:
      - redis_data:/data

volumes:
  redis_data:"""
)

h2("12.4. Pozostale pliki konfiguracyjne")
make_table(
    ["Plik", "Opis"],
    [
        [".python-version", "Wskazuje wersje Pythona (3.13) wymagana przez uv przy tworzeniu srodowiska wirtualnego"],
        [".gitignore", "Wyklucza z repozytorium: __pycache__, pliki .pyc, .venv/, modele (.zip), katalogi best/ i logs/, bazy danych .db, katalog figures/ oraz optuna_study.db"],
        ["uv.lock", "Plik blokady wygenerowany przez uv - precyzyjnie okresla wersje wszystkich zaleznosci (w tym przechodnich) dla powtarzalnosci srodowiska"],
        ["agent/__init__.py, api/__init__.py, state/__init__.py, visualization/__init__.py", "Puste pliki inicjalizujace pakiety Pythona"],
        ["environment/__init__.py", "Eksportuje TrafficLightEnv oraz make_sb3_env jako publiczne API pakietu environment"],
    ],
    col_widths=[6, 10],
)

h2("12.5. models/ - katalog modeli")
para(
    "Katalog models/ zawiera artefakty wytrenowanego agenta. Plik .gitkeep "
    "zapewnia, ze pusty katalog jest sledzony przez git, podczas gdy faktyczne "
    "modele (*.zip) sa wykluczone przez .gitignore i generowane lokalnie podczas "
    "treningu."
)
make_table(
    ["Plik / podkatalog", "Zawartosc"],
    [
        ["ppo_traffic_final.zip", "Finalny model PPO zapisany po zakonczeniu treningu (model.save())"],
        ["best/best_model.zip", "Najlepszy model wedlug EvalCallback (najwyzsza srednia nagroda podczas ewaluacji okresowej)"],
        ["logs/evaluations.npz", "Tablice numpy z historia wynikow ewaluacji okresowej (EvalCallback)"],
    ],
    col_widths=[5, 11],
)

h2("12.6. resources/ - materialy zrodlowe")
make_table(
    ["Plik", "Opis"],
    [
        ["project-requirements.pdf", "Dokument z wymaganiami projektu okreslonymi przez prowadzacego przedmiot"],
        ["documentation.pdf", "Dokumentacja zrodlowa / szablon dokumentacji projektu"],
    ],
    col_widths=[5, 11],
)

page_break()

# =====================================================================
# 13. INSTALACJA I URUCHOMIENIE
# =====================================================================

h1("13. Instalacja i uruchomienie projektu")

h2("13.1. Wymagania")
bullet("Python 3.13 lub nowszy")
bullet("Menedzer pakietow uv (https://docs.astral.sh/uv/)")
bullet("Docker (opcjonalnie, do uruchomienia Redis)")

h2("13.2. Instalacja zaleznosci")
code_block(
"""# Instalacja zaleznosci zgodnie z uv.lock
uv sync

# Uruchomienie Redis (opcjonalnie, do metryk na zywo)
docker compose up -d"""
)

h2("13.3. Trening agenta")
code_block(
"""# Domyslny trening: 300 000 krokow
uv run python scripts/train.py

# Trening niestandardowej dlugosci
uv run python scripts/train.py --timesteps 50000

# Trening bez Redis (jesli Redis nie jest uruchomiony)
uv run python scripts/train.py --timesteps 50000 --no-redis"""
)
para("Wynikiem treningu jest model zapisany w models/ppo_traffic_final.zip, z najlepszym checkpointem w models/best/.")

h2("13.4. Ewaluacja i porownanie z baseline")
code_block(
"""# Ewaluacja headless z porownaniem do baseline
uv run python scripts/evaluate.py --compare-baseline

# Ewaluacja wizualna (renderowanie skrzyzowania)
uv run python scripts/evaluate.py --render --episodes 1"""
)

h2("13.5. Optymalizacja hiperparametrow")
code_block("uv run python scripts/train.py --hyperopt --hyperopt-trials 50")
para("Wyniki zapisywane sa do optuna_study.db i dostepne przez endpoint GET /hyperopt/results.")

h2("13.6. Uruchomienie API")
code_block("uv run python scripts/run_api.py")
para("Interaktywna dokumentacja API (Swagger UI) dostepna jest pod adresem http://localhost:8000/docs.")

h2("13.7. Uruchomienie testow")
code_block("uv run python -m pytest tests/ -v")

page_break()

# =====================================================================
# 14. WYNIKI EKSPERYMENTALNE
# =====================================================================

h1("14. Wyniki eksperymentalne")

para(
    "W ramach realizacji projektu przeprowadzono pelny trening agenta PPO na "
    "domyslnej konfiguracji (300 000 krokow, lambda_a=0.6, lambda_b=0.4, alpha=0.5, "
    "seed=42), a nastepnie wykonano ewaluacje porownawcza z baseline cyklicznym "
    "(cycle_length=15) na 20 epizodach."
)

h2("14.1. Podsumowanie treningu")
make_table(
    ["Parametr przebiegu", "Wartosc"],
    [
        ["Calkowita liczba krokow", "300 000 (303 104 wliczajac ostatni rollout)"],
        ["Czas treningu (CPU)", "ok. 203 s (fps ~1490)"],
        ["Liczba aktualizacji polityki (n_updates)", "730"],
        ["Srednia nagroda przy ewaluacji koncowej (300 000 krokow)", "-30.1 +/- 1.81 (10 epizodow, EvalCallback)"],
        ["Explained variance (ostatnia iteracja)", "0.0409 - 0.0796"],
    ],
    col_widths=[10, 6],
)

h2("14.2. Porownanie agenta PPO z baseline (20 epizodow, seed=42)")
make_table(
    ["Metryka", "Agent PPO", "Baseline (fixed-cycle)"],
    [
        ["Srednia nagroda (mean_reward)", "-30.69", "-66.43"],
        ["Odchylenie standardowe (std_reward)", "+/- 1.71", "+/- 6.52"],
        ["Roznica wzgledem baseline", "+35.74", "-"],
    ],
    col_widths=[6, 5, 5],
)
para(
    "Agent PPO osiaga istotnie wyzsza (mniej ujemna) srednia nagrode oraz znacznie "
    "nizsza wariancje wynikow w porownaniu do baseline. Oznacza to, ze agent nie "
    "tylko skuteczniej minimalizuje sume kolejek, ale rowniez zachowuje sie w "
    "bardziej przewidywalny, stabilny sposob niezaleznie od konkretnego przebiegu "
    "losowych przyjazdow."
)

h2("14.3. Dynamika kolejek - PPO vs baseline")
para(
    "Ponizszy wykres przedstawia przebieg dlugosci kolejek na obu drogach (panele "
    "gorne) oraz harmonogram faz sygnalizacji (panele dolne) w przebiegu jednego "
    "epizodu (seed=42), wygenerowany przez scripts/generate_plots.py."
)
add_figure("queue_dynamics_comparison.png", 16,
           "Rys. 1. Porownanie dynamiki kolejek i harmonogramu faz: agent PPO (lewa kolumna) vs baseline cykliczny (prawa kolumna).")
para(
    "Agent PPO utrzymuje obie kolejki w wezszym, bardziej ograniczonym przedziale "
    "(do ok. 10-12 pojazdow), natomiast baseline cykliczny dopuszcza regularne "
    "wzrosty kolejek do 15-18 pojazdow, poniewaz nie reaguje na rzeczywiste "
    "natezenie ruchu - przelacza sygnalizacje wedlug sztywnego harmonogramu, "
    "niezaleznie od tego, czy danej drodze faktycznie potrzebne jest zielone "
    "swiatlo."
)

h2("14.4. Porownanie sredniej nagrody")
add_figure("reward_comparison.png", 12,
           "Rys. 2. Srednia suma nagrod (obu agentow) w 20 epizodach: agent PPO vs baseline cykliczny, z odchyleniem standardowym.")
para(
    "Wykres slupkowy potwierdza wczesniejsza obserwacje - agent PPO osiaga wyzsza "
    "(mniej ujemna) srednia nagrode przy mniejszym rozrzucie wynikow."
)

h2("14.5. Skumulowana nagroda w czasie")
add_figure("cumulative_reward.png", 14,
           "Rys. 3. Skumulowana nagroda w funkcji kroku epizodu (seed=42): agent PPO vs baseline cykliczny.")
para(
    "Krzywa PPO pozostaje powyzej krzywej baseline przez caly epizod, a roznica "
    "miedzy nimi powieksza sie z czasem - agent PPO skuteczniej zapobiega "
    "narastaniu kolejek, podczas gdy baseline pozwala kolejkom rosnac w sposob "
    "kumulatywny."
)

h2("14.6. Wizualizacja skrzyzowania")
add_figure("intersection_snapshots.png", 17,
           "Rys. 4. Zrzuty wizualizacji skrzyzowania w krokach 1, 100 i 400 epizodu sterowanego przez agenta PPO. "
           "Kolor sygnalizatora oznacza stan: zielony, czerwony lub zolty (faza przejsciowa); paski reprezentuja dlugosc kolejek Q_A i Q_B.")

h2("14.7. Kluczowe wzory matematyczne")

h3("14.7.1. Rozklad Poissona (przyjazdy pojazdow)")
code_block(
"""P(X = k) = (lambda^k * e^(-lambda)) / k!

E[X] = lambda           (oczekiwana liczba przyjazdow na krok)
Var(X) = lambda          (wariancja rowna wartosci oczekiwanej)

Droga A: lambda_A = 0.6 pojazdow/krok
Droga B: lambda_B = 0.4 pojazdow/krok"""
)

h3("14.7.2. Funkcja nagrody")
code_block(
"""scale = max_queue * (1 + alpha) = 50 * 1.5 = 75

reward_A = -(queue_A + alpha * queue_B) / scale
reward_B = -(queue_B + alpha * queue_A) / scale

gdzie alpha = 0.5 (wspolczynnik kooperacji)"""
)
para("Zakres wartosci nagrody: [-1, 0], gdzie 0 oznacza stan optymalny (puste kolejki), a -1 - stan najgorszy (obie kolejki na poziomie pojemnosci max_queue).")

h3("14.7.3. Funkcja celu PPO (clipped objective)")
code_block(
"""L(theta) = E[ min(r(theta) * A, clip(r(theta), 1-eps, 1+eps) * A) ]

gdzie:
  r(theta) = pi_new(a|s) / pi_old(a|s)     (stosunek prawdopodobienstw)
  A = estymata przewagi (advantage)
  eps = 0.2                                 (clip range)"""
)
para("Mechanizm ten zapobiega nadmiernym zmianom polityki w jednej aktualizacji - jezeli r(theta) przekroczyloby 1.2 lub spadlo poniżej 0.8, gradient jest przyciety.")

h3("14.7.4. Generalized Advantage Estimation (GAE)")
code_block(
"""A_t = SUMA_{l=0}^{inf} (gamma * lambda)^l * delta_{t+l}

gdzie delta_t = r_t + gamma * V(s_{t+1}) - V(s_t)    (blad TD)

gamma = 0.99    (wspolczynnik dyskontowania)
lambda = 0.95   (wygladzanie GAE)"""
)

h3("14.7.5. Dynamika kolejki")
code_block(
"""queue_A(t+1) = min(queue_A(t) + arrivals_A(t) - departures_A(t), max_queue)

departures_A(t) = min(queue_A(t), drain_rate)   jezeli droga A jest zielona
                  0                              w przeciwnym przypadku"""
)
para(
    "Stan rownowagi systemu jest osiagany, gdy oczekiwana liczba odjazdow odpowiada "
    "oczekiwanej liczbie przyjazdow. Przy drain_rate=3 oraz lambda_A=0.6 droga A "
    "moze szybko oprozniac swoja kolejke w czasie trwania zielonego swiatla."
)

page_break()

# =====================================================================
# 15. PROBLEMY NAPOTKANE I ROZWIAZANIA
# =====================================================================

h1("15. Problemy napotkane podczas realizacji i ich rozwiazania")

para(
    "W trakcie implementacji i treningu agenta napotkano piec istotnych problemow. "
    "Ich identyfikacja i rozwiazanie mialy kluczowy wplyw na koncowa jakosc "
    "wytrenowanego agenta i sa udokumentowane poniżej jako swiadectwo procesu "
    "inzynierskiego."
)

h2("15.1. Problem 1: brakujaca metoda seed() w SuperSuit")
para("Symptom: AttributeError: 'ConcatVecEnv' object has no attribute 'seed'.")
para(
    "Przyczyna: Stable-Baselines3 wywoluje env.seed() podczas inicjalizacji, a "
    "klasa ConcatVecEnv z SuperSuit nie implementuje tej metody. Wrapper SB3 "
    "delegowal wywolanie seed() do self.venv, dlatego latanie zewnetrznego "
    "wrappera nie wystarczalo."
)
para("Rozwiazanie: monkey-patch wewnetrznego obiektu env.venv:")
code_block(
"""inner = env.venv if hasattr(env, "venv") else env
inner.seed = lambda seed=None: [seed] * env.num_envs"""
)

h2("15.2. Problem 2: agent trwale utrzymuje zielone swiatlo (\"greedy hold\")")
para(
    "Symptom: nagroda PPO nie ulegala poprawie podczas treningu. Droga A "
    "pozostawala zielona przez caly czas, a kolejka na drodze B osiagala maksymalna "
    "wartosc 50."
)
para(
    "Przyczyna: w pierwszej wersji srodowiska tylko agent aktualnie zielony mogl "
    "zainicjowac zmiane fazy. Po uzyskaniu zielonego swiatla agent nie mial zadnej "
    "motywacji, aby je oddac, ponieważ jego wlasna kolejka byla niska. Wspolczynnik "
    "alpha=0.3 nie wystarczal jako kara za zaniedbanie drugiej drogi."
)
para("Rozwiazanie: dopuszczono mozliwosc zglaszania zmiany fazy przez agenta czerwonego oraz zwiekszono alpha z 0.3 do 0.5.")

h2("15.3. Problem 3: funkcja wartosci nie jest w stanie sie nauczyc (skala nagrody za duza)")
para("Symptom: explained_variance pozostawal bliski 0 przez caly trening. Nagroda utrzymywala sie na poziomie ok. -14 000 za epizod.")
para(
    "Przyczyna: surowe nagrody o wartosci ok. -30 na krok generowaly skumulowane "
    "zwroty rzedu -14 000 na epizod. Funkcja wartosci PPO (przewidujaca skumulowana "
    "nagrode przyszla) musiala regresowac wartosci tego rzedu wielkosci - znacznie "
    "trudniejszy problem regresji niz przewidywanie wartosci z waskiego przedzialu."
)
para("Rozwiazanie: normalizacja nagrody do przedzialu [-1, 0]:")
code_block(
"""scale = max_queue * (1.0 + alpha)
reward = -(own_queue + alpha * other_queue) / scale"""
)
para(
    "Metryka explained_variance (0 = przewidywania bezużyteczne, ~1 = przewidywania "
    "blisko idealne) pozwala ocenic, jak dobrze funkcja wartosci przewiduje zwroty - "
    "po normalizacji metryka ta zaczela wzrastac od wartosci bliskich zeru."
)

h2("15.4. Problem 4: w ewaluacji ta sama akcja jest stosowana dla obu agentow")
para(
    "Symptom: w trakcie recznej ewaluacji agent PPO zachowywal sie tak, jakby "
    "byl „zablokowany” - dynamika kolejek nie wykazywala zadnych zmian fazy."
)
para(
    "Przyczyna: kod ewaluacyjny wywolywal model.predict(obs_a) jednokrotnie i "
    "stosowal ten sam wynik akcji do obu agentow. Gdy obaj agenci otrzymywali "
    "akcje 1, aktywowala sie reguła anulowania (cancel-out), blokujac kazda zmiane "
    "fazy."
)
para("Rozwiazanie: niezalezna predykcja dla kazdego agenta:")
code_block(
"""action_a, _ = model.predict(obs["light_A"], deterministic=True)
action_b, _ = model.predict(obs["light_B"], deterministic=True)"""
)

h2("15.5. Problem 5: baseline cykliczny zablokowany przez reguly anulowania")
para("Symptom: baseline cykliczny nigdy nie zmienial fazy sygnalizacji.")
para(
    "Przyczyna: dwie instancje FixedCyclePolicy z identyczna dlugoscia cyklu "
    "uruchamialy sie w tym samym kroku, obie zglaszajac żądanie zmiany - reguła "
    "anulowania blokowala kazda zmiane."
)
para("Rozwiazanie: tylko agent zielony zglasza żądanie zmiany; agent czerwony zawsze zwraca 0:")
code_block(
"""if is_green and self._step % self.cycle_length == 0:
    return 1
return 0"""
)

page_break()

# =====================================================================
# 16. WNIOSKI
# =====================================================================

h1("16. Wnioski i mozliwosci rozwoju")

h2("16.1. Wnioski")
numbered(
    "Agent uczenia ze wzmocnieniem (PPO) jest w stanie nauczyc sie adaptacyjnej "
    "strategii sterowania sygnalizacja swietlna, ktora w sposob mierzalny "
    "przewyzsza klasyczny harmonogram cykliczny - zaobserwowano poprawe sredniej "
    "nagrody o ok. 35 punktow (z -66.43 do -30.69 w skali per-agent) oraz "
    "znaczace zmniejszenie wariancji wynikow (z +/-6.52 do +/-1.71)."
)
numbered(
    "Mechanizm kooperacji (wspolczynnik alpha w funkcji nagrody) jest kluczowy - "
    "bez kary za zaniedbanie kolejki drugiej drogi agenty uczyly sie strategii "
    "„zachlannego utrzymywania” zielonego swiatla, prowadzacej do nieograniczonego "
    "wzrostu kolejki na drugiej drodze."
)
numbered(
    "Projektowanie i normalizacja funkcji nagrody mialy co najmniej tak duzy wplyw "
    "na powodzenie treningu, jak sam wybor algorytmu - nienormalizowane nagrody "
    "uniemozliwialy funkcji wartosci PPO skuteczne uczenie sie."
)
numbered(
    "Symetryczna definicja obserwacji ([wlasna_kolejka, druga_kolejka, faza, czas]) "
    "pozwolila na zastosowanie jednej, wspoldzielonej sieci neuronowej dla obu "
    "agentow (parameter sharing), co jest najprostsza, a jednoczesnie skuteczna "
    "forma uczenia wieloagentowego w tym scenariuszu."
)
numbered(
    "Pelny system (symulacja -> trening -> ewaluacja -> API -> wizualizacja) "
    "zostal zaimplementowany i przetestowany - 26 testow automatycznych przechodzi "
    "poprawnie, a trening 300 000 krokow zajmuje ok. 3.5 minuty na CPU."
)

h2("16.2. Zwiazek z tematyka przedmiotu")
bullet("Proces Poissona: przyjazdy pojazdow modelowane jako zmienne losowe o rozkladzie Poissona - klasyczny proces stochastyczny dyskretny w czasie.")
bullet("Optymalizacja stochastyczna: uczenie ze wzmocnieniem optymalizuje polityke w warunkach niepewnosci (losowe przyjazdy).")
bullet("Szeregi czasowe nagrod: krzywe treningowe przedstawiaja niestacjonarny przebieg nagrody w czasie - typowy obiekt analizy szeregow czasowych.")
bullet("Ocena statystyczna: porownanie srednich i odchylen standardowych nagrod z wielu epizodow jako podstawa wnioskowania o istotnosci poprawy.")

h2("16.3. Mozliwosci dalszego rozwoju")
bullet("Bardziej zlozone skrzyzowania - wiecej drog, wiecej faz sygnalizacji, skrzyzowania wielopoziomowe.")
bullet("Integracja z rzeczywistymi danymi o ruchu drogowym (zamiast syntetycznego procesu Poissona).")
bullet("Koordynacja wielu skrzyzowan (multi-intersection coordination) - rozszerzenie liczby agentow przy zachowaniu wspoldzielonej polityki.")
bullet("Dodanie agentow / ograniczen reprezentujacych pieszych i przejscia dla pieszych.")
bullet("Pelne wykorzystanie wynikow optymalizacji Optuna do automatycznego doboru wspolczynnika kooperacji alpha.")

doc.add_paragraph()
final_p = doc.add_paragraph()
final_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_heading_border(final_p, color="2E75B6", sz=10)
run = final_p.add_run("Koniec dokumentu")
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = BLUE

print("Czesc 10 (problemy + wnioski) wygenerowana - dokument kompletny.")
doc.save(r"C:\Users\Patryk\PycharmProjects\LAB04\LAB13\Traffic_Light_Assistant_Dokumentacja.docx")
